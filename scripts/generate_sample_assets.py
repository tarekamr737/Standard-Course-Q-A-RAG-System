"""Create the small PDF and DOCX samples required for a full format demo.

Run after installing requirements. Generated assets stay under data/ on the D: project
drive and can safely be regenerated.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]


def create_pdf() -> Path:
    output = ROOT / "data" / "course_materials" / "CS4780" / "week_12_handout.pdf"
    output.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(output), pagesize=letter)
    text = pdf.beginText(72, 720)
    text.setFont("Helvetica", 11)
    for line in (
        "CS 4780 Week 12 handout: evaluating regularized models.",
        "Use cross-validation to select the regularization strength.",
        "Do not choose a hyperparameter by repeatedly checking test-set performance.",
        "Ridge uses an L2 penalty. Lasso uses an L1 penalty and can select variables.",
    ):
        text.textLine(line)
    pdf.drawText(text)
    pdf.save()
    return output


def create_docx() -> Path:
    output = ROOT / "data" / "course_materials" / "HIST202" / "week_03_factory_reform.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Week 3 source note: factory reform", level=1)
    document.add_paragraph(
        "The 1833 Factory Act restricted child labor and introduced inspection in British textile factories."
    )
    document.add_paragraph(
        "The course uses this measure to connect industrial production with reform politics and working conditions."
    )
    document.save(output)
    return output


if __name__ == "__main__":
    for path in (create_pdf(), create_docx()):
        print(path.relative_to(ROOT))

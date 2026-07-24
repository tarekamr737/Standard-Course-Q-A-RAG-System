"""Safe, metadata-preserving loaders for supported course-material formats."""

from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from .models import Document


SUPPORTED_EXTENSIONS = {".pdf", ".csv", ".docx", ".txt"}


class UnsupportedFormatError(ValueError):
    pass


class MalformedDocumentError(ValueError):
    pass


def clean_text(value: str) -> str:
    """Normalize whitespace without removing useful sentence boundaries."""
    value = value.replace("\x00", " ").replace("\u00a0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _metadata(course: str, path: Path, **location: str | int) -> dict[str, str | int]:
    return {
        "course": course,
        "file_name": path.name,
        "file_type": path.suffix.removeprefix(".").upper(),
        **location,
    }


def load_txt(path: Path, course: str) -> list[Document]:
    try:
        text = clean_text(path.read_text(encoding="utf-8-sig", errors="replace"))
    except OSError as error:
        raise MalformedDocumentError(f"Could not read {path.name}.") from error
    if not text:
        raise MalformedDocumentError(f"{path.name} does not contain readable text.")
    sections = [clean_text(section) for section in re.split(r"\n(?=#{1,3}\s)|\n(?=[A-Z][A-Z ]{4,}:)", text)]
    documents: list[Document] = []
    pending_heading = ""
    for index, section in enumerate(section for section in sections if section):
        if section.startswith("#") and "\n\n" not in section and len(section) < 160:
            pending_heading = section
            continue
        combined = "\n\n".join(part for part in (pending_heading, section) if part)
        combined = re.sub(r"(?m)^#{1,6}\s*", "", combined)
        documents.append(Document(clean_text(combined), _metadata(course, path, section=index + 1)))
        pending_heading = ""
    if pending_heading:
        documents.append(Document(re.sub(r"^#{1,6}\s*", "", pending_heading), _metadata(course, path, section=len(sections))))
    if not documents:
        raise MalformedDocumentError(f"{path.name} does not contain readable sections.")
    return documents


def load_csv(path: Path, course: str) -> list[Document]:
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.DictReader(handle))
    except (OSError, csv.Error, UnicodeError) as error:
        raise MalformedDocumentError(f"Could not parse CSV file {path.name}.") from error
    if not rows or not rows[0]:
        raise MalformedDocumentError(f"{path.name} does not contain rows with headers.")
    documents = []
    for row_number, row in enumerate(rows, start=2):
        text = clean_text(". ".join(f"{key}: {value}" for key, value in row.items() if value))
        if text:
            documents.append(Document(text, _metadata(course, path, row=row_number)))
    if not documents:
        raise MalformedDocumentError(f"{path.name} has no non-empty rows.")
    return documents


def load_pdf(path: Path, course: str) -> list[Document]:
    try:
        reader = PdfReader(str(path))
        documents = [
            Document(clean_text(page.extract_text() or ""), _metadata(course, path, page=index + 1))
            for index, page in enumerate(reader.pages)
            if clean_text(page.extract_text() or "")
        ]
    except Exception as error:  # pypdf exposes several parser exceptions.
        raise MalformedDocumentError(f"Could not extract readable PDF text from {path.name}.") from error
    if not documents:
        raise MalformedDocumentError(f"{path.name} has no extractable text.")
    return documents


def load_docx(path: Path, course: str) -> list[Document]:
    try:
        document = DocxDocument(path)
        paragraphs = [clean_text(paragraph.text) for paragraph in document.paragraphs]
    except Exception as error:
        raise MalformedDocumentError(f"Could not parse DOCX file {path.name}.") from error
    documents = [
        Document(text, _metadata(course, path, section=index + 1))
        for index, text in enumerate(paragraphs)
        if text
    ]
    if not documents:
        raise MalformedDocumentError(f"{path.name} has no readable paragraphs.")
    return documents


def load_file(path: Path, course: str) -> list[Document]:
    """Load one supported document, retaining file and location metadata."""
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFormatError(f"{path.name} is unsupported. Use one of: {supported}.")
    loaders = {".txt": load_txt, ".csv": load_csv, ".pdf": load_pdf, ".docx": load_docx}
    return loaders[suffix](path, course)


def discover_course_files(course: str, directories: list[Path]) -> list[Path]:
    files: list[Path] = []
    for directory in directories:
        course_directory = directory / course
        if course_directory.exists():
            files.extend(path for path in course_directory.rglob("*") if path.suffix.lower() in SUPPORTED_EXTENSIONS)
    return sorted(set(files))
